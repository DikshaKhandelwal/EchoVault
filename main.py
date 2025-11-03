import io
import os
import time
from typing import Any, Dict, List

import numpy as np
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from ai_utils import cosine_sim, embed_text, summarize_text
from storage import get_all, get_by_path, update_last_used, upsert_file
from watcher import diff_with_db, scan_folder, suggest_archive

# Load environment variables
load_dotenv()

app = FastAPI(title="EchoVault API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class IngestFolderReq(BaseModel):
    folder: str


class RecallReq(BaseModel):
    query: str
    top_k: int = 5
    decay_days: int = 120  # memory aging


def extract_text_from_ext(path: str, content: bytes | None = None) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in {".txt", ".md"}:
        data = content if content is not None else open(path, "rb").read()
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return data.decode("latin-1", errors="ignore")
    if ext == ".pdf":
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content) if content is not None else path)
        out = []
        for page in reader.pages:
            out.append(page.extract_text() or "")
        return "\n".join(out)
    if ext == ".docx":
        import docx

        document = docx.Document(io.BytesIO(content)) if content is not None else docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs)
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")


@app.post("/ingest/folder")
def ingest_folder(req: IngestFolderReq) -> Dict[str, Any]:
    folder = os.path.abspath(req.folder)
    if not os.path.isdir(folder):
        raise HTTPException(400, detail="Folder not found")
    added: List[str] = []
    updated: List[str] = []
    for path, mtime in scan_folder(folder):
        with open(path, "rb") as f:
            text = extract_text_from_ext(path, f.read())
        summary, tags = summarize_text(text)
        emb = embed_text(text)
        rec = get_by_path(path)
        upsert_file(path, summary, emb, mtime, tags)
        (updated if rec else added).append(path)
    return {"added": added, "updated": updated}


@app.post("/ingest/upload")
def ingest_upload(file: UploadFile = File(...)) -> Dict[str, Any]:
    content = file.file.read()
    try:
        text = extract_text_from_ext(file.filename, content)
    finally:
        file.file.close()
    mtime = time.time()
    summary, tags = summarize_text(text)
    emb = embed_text(text)
    upsert_file(f"uploaded://{file.filename}", summary, emb, mtime, tags)
    return {"path": f"uploaded://{file.filename}", "tags": tags}


@app.post("/recall")
def recall(req: RecallReq) -> Dict[str, Any]:
    query_emb = embed_text(req.query)
    items = get_all()
    now = time.time()
    scored = []
    for r in items:
        sim = cosine_sim(query_emb, r["embedding"]) if r["embedding"] else 0.0
        # memory aging
        age_days = (now - (r["last_used"] or r["last_modified"] or now)) / 86400.0
        decay = max(0.5, 1.0 - (age_days / max(1, req.decay_days)))
        score = sim * decay
        scored.append((score, r))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [
        {
            "path": r["path"],
            "summary": r["summary"],
            "tags": r["tags"],
            "score": round(s, 4),
        }
        for s, r in scored[: req.top_k]
    ]
    for t in top:
        update_last_used(t["path"])
    return {"results": top}


@app.post("/chat")
def chat(req: RecallReq) -> Dict[str, Any]:
    """RAG-based chat: answer questions using file contents"""
    from ai_utils import chat_with_context
    
    # Get top relevant files
    query_emb = embed_text(req.query)
    items = get_all()
    scored = []
    for r in items:
        sim = cosine_sim(query_emb, r["embedding"]) if r["embedding"] else 0.0
        scored.append((sim, r))
    scored.sort(key=lambda x: x[0], reverse=True)
    
    # Build contexts for RAG
    contexts = []
    for _, r in scored[:5]:
        fname = os.path.basename(r["path"])
        summary = r.get("summary", "")
        content = r.get("content", "")[:1000]  # snippet
        contexts.append((fname, summary, content))
    
    if not contexts:
        return {"answer": "No files found to answer this question."}
    
    answer = chat_with_context(req.query, contexts)
    return {"answer": answer, "sources": [c[0] for c in contexts]}


@app.get("/sync")
def sync(folder: str, archive_days: int = 60) -> Dict[str, Any]:
    result = diff_with_db(folder)
    archives = suggest_archive(days=archive_days)
    return {"diff": result, "archive_suggestions": archives}


@app.get("/files")
def list_files() -> Dict[str, Any]:
    return {"files": get_all()}


@app.get("/demo/age-files")
def age_demo_files(pattern: str = "old_projects", days: int = 90):
    from watcher import age_files_for_demo
    age_files_for_demo(pattern, days)
    return {"status": "ok", "aged": pattern, "days": days}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")))

